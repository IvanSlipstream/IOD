from StringIO import StringIO
from datetime import datetime, timedelta

from django.http import HttpResponse
from django.template.loader import render_to_string
from docx import Document

from RFC import constants
from RFC import tools
from RFC.models import ChangeRequest, UserMeta

class RFCDocument():
    def __init__(self, file):
        self.document = Document(file)

    def set_body(self, body):
        table = self.document.tables[0]
        cell = table.cell(3, 0)
        cell.text = body

    def set_oper_our(self, oper_our):
        table = self.document.tables[0]
        cell = table.cell(1, 0)
        paragraph = cell.paragraphs[1]
        paragraph.text = oper_our

    def set_author(self, author):
        table = self.document.tables[0]
        cell = table.cell(1, 2)
        paragraph = cell.paragraphs[0]
        paragraph.text = author

    def set_time(self, time):
        table = self.document.tables[0]
        cell = table.cell(1, 5)
        cell.text = time

    def set_responsible(self, responsible, approved):
        table = self.document.tables[1]
        cell = table.cell(0, 1)
        cell.text = responsible
        cell = table.cell(2, 1)
        cell.text = approved

    def set_number(self, number):
        paragraph = self.document.paragraphs[2]
        paragraph.text += str(number)

    def set_comments(self, comments):
        table = self.document.tables[0]
        cell = table.cell(5, 4)
        cell.text = comments

    def save(self, file):
        self.document.save(file)


def inflate_docx(rfc_base_id, chosen_rfcs=None):
    if chosen_rfcs is None:
        chosen_rfcs = []
    if rfc_base_id not in chosen_rfcs:
        chosen_rfcs.append(rfc_base_id)
    rfc_number_ref = tools.adopt(chosen_rfcs)
    rfc_base = ChangeRequest.objects.get(id=rfc_base_id)
    rfcs = ChangeRequest.objects.filter(id__in=chosen_rfcs)
    rfcs_forward = rfcs.filter(direction=1)
    rfcs_backward = rfcs.filter(direction=2)
    rfcs_two_way = rfcs.filter(direction=3)
    c = {'oper_our': rfc_base.oper_our.fineName,
         'rfcs_forward': rfcs_forward, 'rfcs_backward': rfcs_backward, 'rfcs_two_way': rfcs_two_way}
    body = render_to_string('body.html', dictionary=c)
    out = StringIO()
    document = RFCDocument("C:\\Users\\Slipstream\\Desktop\\change_request.docx")
    document.set_body(body)
    document.set_oper_our(rfc_base.oper_our.fineName)
    document.set_number(rfc_number_ref)
    document.set_author(UserMeta.objects.get(user_reference=rfc_base.author).__str__().decode('utf-8'))
    document.set_time((datetime.now().replace(minute=0, second=0) + timedelta(hours=1))
                      .strftime('%Y-%m-%d\n%H:%M:%S'))
    document.set_comments(rfc_base.comments)
    document.set_responsible(constants.OU_IW_HEAD, constants.TD_HEAD)
    document.save(out)
    response = HttpResponse(out.getvalue(), mimetype='application/docx')
    response['Content-Disposition'] = 'attachment; filename=change_request_%s.docx' \
                                      % (rfc_base.oper_our.fineName.replace(" ", "_"))
    out.close()
    return response
